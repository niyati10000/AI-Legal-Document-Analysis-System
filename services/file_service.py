import os
import pypdf
import docx
import traceback

def extract_text_from_file(filepath):
    """
    Unified text extraction service for PDF, DOCX, and TXT files.
    Returns the extracted text as a string.
    Raises exceptions if the file type is unsupported or extraction fails.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found at path: {filepath}")

    filename = os.path.basename(filepath)
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext == 'txt':
        return extract_from_txt(filepath)
    elif ext == 'pdf':
        return extract_from_pdf(filepath)
    elif ext == 'docx':
        return extract_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Supported formats are: PDF, DOCX, TXT.")

def extract_from_txt(filepath):
    """Extract text from plain text files with encoding fallback"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
                # Clean up carriage returns
                return content.replace('\r\n', '\n').strip()
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("txt", b"", 0, 0, "Failed to decode text file with standard encodings.")

def extract_from_pdf(filepath):
    """Extract text from PDF documents using pypdf"""
    try:
        reader = pypdf.PdfReader(filepath)
        text_pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_pages.append(page_text)
            else:
                text_pages.append(f"\n--- [Page {i+1} Empty or Scanned Image] ---\n")
        
        extracted_text = "\n".join(text_pages).strip()
        if not extracted_text.replace("\n", "").strip():
            raise ValueError("No text could be extracted from the PDF. It might be scanned/image-only.")
        return extracted_text
    except Exception as e:
        print(f"Error parsing PDF file {filepath}: {e}")
        traceback.print_exc()
        raise RuntimeError(f"Failed to parse PDF file: {str(e)}")

def extract_from_docx(filepath):
    """Extract text from MS Word documents using python-docx"""
    try:
        doc = docx.Document(filepath)
        paragraphs = []
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    paragraphs.append(" | ".join(row_cells))
        
        extracted_text = "\n\n".join(paragraphs).strip()
        if not extracted_text:
            raise ValueError("Word document appears to be empty.")
        return extracted_text
    except Exception as e:
        print(f"Error parsing Word file {filepath}: {e}")
        traceback.print_exc()
        raise RuntimeError(f"Failed to parse Word document: {str(e)}")

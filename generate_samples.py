import os
import docx

os.makedirs('sample_documents', exist_ok=True)

# 1. Create a sample NDA in text format (contains PII for masking test)
nda_text = """MUTUAL NON-DISCLOSURE AGREEMENT

This Mutual Non-Disclosure Agreement ("Agreement") is entered into on August 14, 2026 ("Effective Date"), by and between:
Disclosing Party: Sarah Jenkins (Email: sarah.j@techcorp.com, Phone: 212-555-0199, SSN: 000-12-3456)
Receiving Party: Michael Vance (Email: m.vance@innovatelabs.org, Phone: 415-555-0245)

1. Purpose
The parties wish to explore a business relationship of mutual interest. In connection with this, the Disclosing Party may disclose proprietary information.

2. Confidential Information
Confidential Information refers to any proprietary information, technical data, trade secrets, or know-how, including but not limited to research, product plans, source code, and software, disclosed by the Disclosing Party.

3. Redaction of Personal Data
The parties agree that any personal identifiers, including social security numbers, banking credit cards (e.g. 4111-2222-3333-4444), and private contact details shall be handled in accordance with GDPR and compliance regulations.

4. Arbitration Clause
Any dispute arising out of or in connection with this contract shall be determined by arbitration in New York City, before a single arbitrator.

5. Term
This Agreement and the Receiving Party's duty to hold Confidential Information in confidence shall remain in effect for three (3) years from the Effective Date.

IN WITNESS WHEREOF, the parties have executed this Agreement.

__________________________
Sarah Jenkins
Disclosing Party

__________________________
Michael Vance
Receiving Party
"""

with open('sample_documents/sample_nda.txt', 'w', encoding='utf-8') as f:
    f.write(nda_text)

# 2. Create a sample Employment Agreement in DOCX format (contains biased terms for auditing tests)
doc = docx.Document()
doc.add_heading('EMPLOYMENT AGREEMENT', 0)

doc.add_paragraph('This Employment Agreement ("Agreement") is made effective as of August 1, 2026, by and between Zenith Legal Associates LLC ("Employer") and Jane Doe ("Employee").')

p1 = doc.add_paragraph()
p1.add_run('1. Duties and Execution. ').bold = True
p1.add_run('The Employee is hired as a Senior Counsel. We expect the Employee to be highly committed. Historically, female attorneys in this role have struggled with childcare duties. Therefore, we require the Employee to demonstrate aggressive and unyielding commitment, avoiding strident or emotional reactions during negotiations. We seek a fresh graduate who is young, adaptable, and willing to work extended hours without family-related interruptions.')

p2 = doc.add_paragraph()
p2.add_run('2. Compensation. ').bold = True
p2.add_run('Zenith Legal agrees to pay the Employee an annual salary of $120,000, payable in semi-monthly installments.')

p3 = doc.add_paragraph()
p3.add_run('3. Non-Compete. ').bold = True
p3.add_run('The Employee agrees that during the term of employment and for a period of one (1) year following termination, they shall not engage in any competitive legal services within a 50-mile radius of the Employer\'s office.')

doc.save('sample_documents/sample_employment_contract.docx')
print("Sample documents generated successfully!")
